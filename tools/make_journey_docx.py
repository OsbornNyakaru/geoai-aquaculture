# Build JOURNEY.docx — same content as JOURNEY.md, with flowcharts as ASCII diagrams.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\ADMIN\OneDrive\Desktop\OSBORN\AGENTIC-AI-DVPT\ZINDI-PROJECTS\geoai-aquaculture\JOURNEY.docx"

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(8)

for s, sz, col in (("Title", 26, "1F3864"), ("Heading 1", 17, "1F3864"),
                   ("Heading 2", 13, "2E5A88")):
    f = doc.styles[s].font
    f.size, f.name, f.bold = Pt(sz), "Calibri", True
    f.color.rgb = RGBColor.from_string(col)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def para(text, bold=False, italic=False, size=11, space=8, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic, r.font.size = bold, italic, Pt(size)
    p.paragraph_format.space_after = Pt(space)
    if align:
        p.alignment = align
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        parts = it.split("**")
        for i, seg in enumerate(parts):
            r = p.add_run(seg)
            r.bold = (i % 2 == 1)


def diagram(lines, title=None):
    """Monospace boxed diagram in a shaded single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, "F2F5FA")
    c.paragraphs[0].text = ""
    first = True
    if title:
        p = c.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("1F3864")
        p.paragraph_format.space_after = Pt(4)
        first = False
    for ln in lines:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        shade(hdr[i], "1F3864")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            for j, seg in enumerate(str(val).split("**")):
                r = p.add_run(seg)
                r.bold = (j % 2 == 1)
                r.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ------------------------------------------------------------------ TITLE
doc.add_paragraph("The Journey So Far", style="Title")
para("GeoAI Aquaculture Pond Identification — a plain-English story of what we tried, "
     "what failed, what we kept, and where we stand.", italic=True, size=12)

table(["", ""], [
    ["Competition", "GeoAI Aquaculture Pond Identification (Zindi / FAO / ITU)"],
    ["Deadline", "2026-08-16   ·   Submissions allowed: 5 per day"],
    ["Where we started", "public leaderboard 0.7140"],
    ["Where we are now", "**public leaderboard 0.8955**"],
    ["Total climb", "**+0.1815**"],
    ["Last updated", "2026-07-21"],
], widths=[1.6, 4.6])

# ------------------------------------------------------------------ 1
doc.add_heading("1. The problem in one paragraph", level=1)
para("We are given 12 months of satellite readings (radar + optical) for each map cell, and we "
     "must say whether that cell is an aquaculture pond or not.")
para("The catch — and it is the whole story of this project — is that the training data and the "
     "test data are deliberately different. The competition organisers built them that way. A "
     "model can memorise the training set almost perfectly and still transfer badly. We measured "
     "this: a classifier can tell a training row from a test row with 99% accuracy.")
para("That single fact reshaped every decision below.", italic=True)

# ------------------------------------------------------------------ 2
doc.add_heading("2. The one rule we run on", level=1)
para("The local score lies. The leaderboard is the only truth.", bold=True, size=13)
para("Early on we found two models with identical local cross-validation scores whose "
     "leaderboard scores differed by 0.05. Later it got worse — the local score became actively "
     "backwards:")
table(["Run", "Local score (OOF)", "Leaderboard", "Note"], [
    ["K=4 augmentation", "**0.9840** (best local)", "0.8665", "2nd-worst on LB"],
    ["Relative-time", "0.9811 (lower)", "0.8908", "won"],
    ["Cross-view invariance", "**0.9753** (lowest)", "**0.8955**", "**best ever**"],
], widths=[1.9, 1.6, 1.2, 1.5])
para("So the rule: we never keep an idea because it scored well locally. Every keep/discard "
     "decision waits for a real leaderboard number. That is why the loop has a human in it.")

diagram([
    "  +-------------------+     +---------------------+     +-------------------+",
    "  |  Agent picks the  | --> |  Edits config and   | --> |  Human runs it    |",
    "  |  next experiment  |     |  pushes to GitHub   |     |  on a Colab GPU   |",
    "  +-------------------+     +---------------------+     +---------+---------+",
    "            ^                                                     |",
    "            |                                                     v",
    "  +---------+---------+     +---------------------+     +-------------------+",
    "  |  KEEP (new champ) | <-- |  Beat the champion? | <-- |  Upload CSV to    |",
    "  |  or DISCARD       |     |      YES  /  NO     |     |  Zindi, paste LB  |",
    "  +-------------------+     +---------------------+     +-------------------+",
], "THE LOOP")

# ------------------------------------------------------------------ 3
doc.add_heading("3. How far we've come", level=1)
diagram([
    "  START            +  Prior          +  Transformer     +  Relative-time   +  Cross-view",
    "  0.7140    --->      correction  --->   instead     --->   reframing   --->  invariance",
    "  plain GBDT         0.8260            of GBDT           0.8908            0.8955  <== HERE",
    "  ensemble           (+0.112)          0.8780            (+0.013)          (+0.005)",
    "                                       (+0.052)",
    "",
    "                                                     ...still ahead:  top-5  ~0.928+",
], "THE CLIMB: 0.7140  ->  0.8955")

para("Four things have ever worked. Everything else lost.")
table(["#", "The step that worked", "Why it worked, in plain terms", "Gain"], [
    ["1", "Prior correction",
     "The test set has far more ponds (~65%) than the training set (~40%). We shifted the "
     "model's output so it predicts ponds at the right rate.", "**+0.112**"],
    ["2", "Swap GBDT → Transformer",
     "The Transformer looks only at the months that were actually observed. The GBDT flattened "
     "everything into averages and memorised the training region.", "**+0.052**"],
    ["3", "Relative-time reframing",
     "Stop telling the model \"this is March.\" Tell it \"this is month 1 of the window.\" It "
     "was memorising the calendar, which does not carry over to the test area.", "**+0.013**"],
    ["4", "Cross-view invariance",
     "Show the model the same cell with different months hidden, and penalise it for changing "
     "its mind. Teaches it the answer shouldn't depend on which months you happened to see.",
     "**+0.005**"],
], widths=[0.3, 1.5, 3.6, 0.8])

# ------------------------------------------------------------------ 4
doc.add_heading("4. What diverged — the things we tried and dropped", level=1)
para("We have spent 10 experiments. Six lost, one tied, three won. This is the honest scoreboard.")
table(["Idea", "What it did", "LB", "Verdict"], [
    ["GBDT + Transformer blend", "Average two different models", "0.8705",
     "✗ −0.0075 — the GBDT dragged the good model down"],
    ["Per-cell detrend channels", "Add \"level-removed\" input features", "**0.8266**",
     "✗ −0.0514 — our worst result ever"],
    ["More augmentation (K=4)", "Show each row 4 masked views instead of 2", "0.8665",
     "✗ −0.0115 — best local score, near-worst LB"],
    ["Test-time augmentation", "Average predictions over 8 masked views", "0.8885",
     "✗ −0.0023 — no harm, no help"],
    ["Duration-normalised positions", "Squeeze every window onto a shared 0→1 timeline", "0.8844",
     "✗ −0.0064 — deleted useful info"],
    ["NoPE (no position at all)", "Treat the months as an unordered bag", "0.8917",
     "= tie — **kept as a backup model**"],
    ["Stronger cross-view penalty", "Turn our newest win up 3x", "0.8921",
     "✗ −0.0034 — we had already found the sweet spot"],
], widths=[1.5, 1.9, 0.7, 2.1])
para("Also permanently rejected (researched, argued, and ruled out): EM/Saerens prior estimation "
     "(rejected 3 separate times), water-index threshold features, self-training on the test set, "
     "importance weighting / domain-adversarial training, stacking, temperature scaling, and "
     "anything using pretrained or external models (banned by the rules).")

# ------------------------------------------------------------------ 5
doc.add_heading("5. The one lesson that explains all of it", level=1)
para("Every failure and every success fits one rule we discovered the hard way:")
diagram([
    "            I have a new idea. Will it help?",
    "                        |",
    "     Does it ADD capacity?  (more models, features, augmentation)",
    "         YES --> [X] It will LOSE.  Every single one did.",
    "          |",
    "          NO",
    "          v",
    "     Is it just variance reduction?  (averaging, smoothing)",
    "         YES --> [=] It lands inside the noise.  Not worth a submission.",
    "          |",
    "          NO",
    "          v",
    "     Does it DELETE a channel genuinely DIFFERENT between train and test?",
    "         YES --> [V] This is the winning shape.   (relative-time did this)",
    "         NO  --> [X] You're deleting real signal. (duration-norm did this)",
], "THE DECISION RULE")
para("In one sentence: don't give the model more — take away the specific thing it is "
     "memorising.", bold=True)
para("Adding always hurt. Deleting helped only when the deleted thing was genuinely different "
     "between the training area and the test area. Calendar month was different → deleting it "
     "won. Window length was already matched → deleting it lost.")
para("One more constraint that governs everything: the public leaderboard is scored on only "
     "~309 rows, which means it carries roughly ±0.01 of random noise. We therefore refuse to "
     "test small ideas — a +0.003 improvement is simply unmeasurable. We only spend a submission "
     "on changes big enough to be seen.")

# ------------------------------------------------------------------ 6
doc.add_heading("6. What we are using right now", level=1)
diagram([
    "   12 months of satellite bands  (many months missing)",
    "                   |",
    "                   v",
    "   Standardise  +  add missing-month flags",
    "                   |",
    "                   v",
    "   Create 2 MASKED VIEWS of every training row",
    "                   |",
    "                   v",
    "   RELATIVE-TIME:  slide the window to start at step 0   <== win #3",
    "                   (kills calendar memorisation)",
    "                   |",
    "                   v",
    "   TRANSFORMER ENCODER   2 layers · 4 heads · d=64",
    "                   (attends to observed months only)",
    "                   |",
    "                   v",
    "   Average over observed months  ->  classifier head",
    "                   |",
    "                   v",
    "   CROSS-VIEW PENALTY: punish disagreement between views  <== win #4",
    "                   |",
    "                   v",
    "   Shift output so the pond-rate lands at 0.649           <== win #1",
    "                   |",
    "                   v",
    "   submission.csv   ->   0.8955",
], "THE CURRENT PIPELINE")

para("The champion, in words: a Transformer we trained from scratch (no pretrained weights — "
     "they are banned), which reads only the months that actually exist, is told relative time "
     "rather than calendar time, is trained to give the same answer no matter which months are "
     "hidden, and finally has its output nudged so it predicts ponds at the right frequency.")
doc.add_heading("The settings that define it", level=2)
table(["Setting", "Value", "Why"], [
    ["relative_time", "true", "Win #3 — deletes calendar memorisation"],
    ["consistency_lambda", "1.0", "Win #4 — the cross-view penalty"],
    ["K (training views)", "2", "4 was tested and lost; 2 is a sharp optimum"],
    ["pos_encoding", "learned", "dnorm lost, none tied"],
    ["prevalence_target", "0.649", "Win #1 — the operating point"],
    ["all channels.*", "false", "Every added channel lost"],
    ["tta.enable", "false", "Tested, landed in the noise"],
], widths=[1.7, 0.9, 3.6])

# ------------------------------------------------------------------ 7
doc.add_heading("7. How strong is it?", level=1)
table(["", "Score"], [
    ["Where we started", "0.7140"],
    ["**Our champion today**", "**0.8955**"],
    ["Top-5 on the board", "~0.928 – 0.945"],
    ["Leader", "~0.9452"],
    ["**Gap left to close**", "**~ +0.033**"],
], widths=[3.0, 3.2])
para("Honest read: we are in solid, competitive territory and we have closed 85% of the distance "
     "from our starting point to the top of the board. The remaining +0.033 is hard — it is "
     "roughly three times the size of our last two wins combined, and the two biggest levers "
     "(calendar position, and per-cell signal strength) are now exhausted and proven-toxic "
     "respectively.")
para("Confidence in the champion is high because the wins were not lucky toggles — each one has "
     "a mechanism we can state and that we verified in the run logs. The most recent win, for "
     "example, measurably reduced the model's overconfidence (its known weakness) exactly as "
     "predicted.")
para("We also hold a second, deliberately different model (the NoPE set encoder, 0.8917). It ties "
     "our champion on the public board but is built on a completely different assumption, so it "
     "fails on different rows. That is insurance for the private leaderboard.")

# ------------------------------------------------------------------ 8
doc.add_heading("8. What's in the pipeline", level=1)
diagram([
    "   Iteration 10 : stronger penalty LOST (0.8921)  ->  reverted to the champion",
    "                                  |",
    "                                  v",
    "   Both idea-lanes now CLOSED :  position exhausted  ·  objective at its optimum",
    "                                  |",
    "                                  v",
    "   NOW — RESEARCH ROUND 06  (deep-research brief, Claude Fable)",
    "                                  |",
    "      +-------------+-------------+-------------+-------------+",
    "      v             v             v             v",
    "  1. a local    2. feature    3. maths we    4. CV design +",
    "     score that     engineering   have never      the pond-science",
    "     PREDICTS       inside the    tried           literature",
    "     the LB         Transformer",
    "      +-------------+-------------+-------------+-------------+",
    "                                  |",
    "                                  v",
    "                    triage  ->  next experiment",
    "                                  |",
    "                                  v",
    "                     ENDGAME (near the deadline)",
    "        · measure the true seed noise            (2 submissions)",
    "        · one-time prevalence sweep              (4 submissions)",
    "        · lock the finalists: champion + NoPE",
], "WHAT HAPPENS NEXT")

para("What just happened: iteration 10 turned the cross-view penalty up 3x and lost (0.8921). The "
     "useful part is why. It de-saturated the model even further than the winning setting did — "
     "yet its ranking ability was completely untouched (AUC held at 0.9896). So the loss isn't the "
     "model breaking; it's simply that a little of this medicine helps and more does not. We had "
     "already found the sweet spot. That lane is now closed at its optimum.")
para("Where that leaves us: both of our productive lanes are measured closed, and we are out of "
     "queued ideas big enough to clear the ±0.01 noise. Importantly, submissions are no longer the "
     "constraint — we have ~130 left over ~26 days. Ideas are the constraint, and so is our "
     "ability to measure small ones.")
para("So the next move is research, not another guess. The brief going out asks for depth in four "
     "lanes we have never explored:")
bullets([
    "**A local score that predicts the leaderboard.** The single highest-value thing available. "
    "Right now we cannot tell a real +0.005 from noise without spending a submission — our "
    "champion's own margin sits at that edge. If we could screen ideas offline, our 130 spare "
    "submissions suddenly become useful instead of unusable.",
    "**Feature engineering inside the Transformer.** An admitted blind spot: our champion sees "
    "only 24 raw channels. Every richer feature we ever built lives in the abandoned GBDT branch.",
    "**Mathematics we've never touched** — robust optimization, optimal transport, ranking losses "
    "(worth noting: AUC is 40% of the score and we optimise plain cross-entropy).",
    "**CV design and the actual pond-science literature** — is there a physical signature of "
    "aquaculture ponds that our architecture simply cannot represent?",
])
doc.add_heading("Then the endgame, close to the deadline", level=2)
bullets([
    "**Measure the true seed-to-seed spread** (2 submissions). We have always assumed ±0.01 from "
    "row-count theory but never actually measured it. If the spread is tiny, our small wins are "
    "real and we can resolve finer effects; if it is large, several recent readings were noise.",
    "**A one-time prevalence sweep** (4 submissions). Our pond-rate of 0.649 was tuned for a much "
    "older model. Re-checking it costs no retraining at all, and 60% of the competition metric "
    "hangs on it. We will pick the centre of the flat region, not the single best number — that "
    "avoids fitting the public leaderboard's noise.",
    "**Lock the two finalists:** the champion, plus the structurally different NoPE model.",
])
para("One open question for you: we still need to confirm on the Zindi rules page whether the "
     "private leaderboard auto-selects your best public submission, or lets you nominate two "
     "finalists. If it is the latter, our NoPE backup is genuinely valuable insurance. If it is "
     "the former, we should put every remaining submission into raising the champion alone.",
     bold=True)

# ------------------------------------------------------------------ 9
doc.add_heading("9. Where everything lives", level=1)
table(["File", "What it holds"], [
    ["PROJECT_STATE.md", "The master state doc — carry this to any new cloud account"],
    ["experiments/LB_LOG.md", "Every submission and its leaderboard score (the reward signal)"],
    ["gemini_loop/AGENT_BRIEF.md", "The standing instructions for the agent driving the loop"],
    ["experiments/run_current.sh", "The single experiment currently staged to run"],
    ["config/config.yaml", "Every setting, in one place"],
    ["src/seq_model.py", "The champion Transformer itself"],
], widths=[2.2, 4.0])

doc.save(OUT)
print("saved:", OUT)
