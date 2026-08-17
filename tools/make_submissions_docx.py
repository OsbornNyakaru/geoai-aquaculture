"""Build SUBMISSIONS_LEDGER.docx -- the full Zindi submission record, downloadable/printable.

Same house style as `tools/make_journey_docx.py`. Landscape, because the ledger is 10 columns wide.

Contents
  1. headline result + how to read the tables
  2. all 91 submissions: id, date, FILENAME, finalist flag, public/private composite, AUC, F1
  3. the final leaderboard top 25, plus our row, with the same decomposition
  4. the peer comparison at matched AUC that explains the gap

Sources are the two TSVs written when the competition closed:
  experiments/zindi_submissions_final.tsv
  experiments/zindi_final_leaderboard_top75.tsv

USAGE
    python tools/make_submissions_docx.py
"""
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "SUBMISSIONS_LEDGER.docx"

LEDGER = ROOT / "experiments" / "zindi_submissions_final.tsv"
BOARD = ROOT / "experiments" / "zindi_final_leaderboard_top75.tsv"
COLS = ["id", "when_raw", "file", "sel", "pub", "prv", "auc_pub", "auc_prv", "f1_pub", "f1_prv"]

NAVY, BLUE, GREY = "1F3864", "2E5A88", "595959"
HDR_FILL, ALT_FILL, HILITE = "1F3864", "F2F5FA", "FFF2CC"

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Inches(0.5)
sec.top_margin = sec.bottom_margin = Inches(0.5)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)
st.paragraph_format.space_after = Pt(6)
for s, sz, col in (("Title", 24, NAVY), ("Heading 1", 15, NAVY), ("Heading 2", 12, BLUE)):
    f = doc.styles[s].font
    f.size, f.name, f.bold = Pt(sz), "Calibri", True
    f.color.rgb = RGBColor.from_string(col)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def para(text, bold=False, italic=False, size=10, space=6, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic, r.font.size = bold, italic, Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(space)
    return p


def table(headers, rows, widths=None, highlight=None, mono_cols=()):
    """highlight: predicate(row_index, row) -> bool, shades the row."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, HDR_FILL)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold, r.font.size = True, Pt(8)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        hot = highlight(ri, row) if highlight else False
        for ci, val in enumerate(row):
            c = cells[ci]
            if hot:
                shade(c, HILITE)
            elif ri % 2 == 1:
                shade(c, ALT_FILL)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(8)
            r.bold = hot
            if ci in mono_cols:
                r.font.name = "Consolas"
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# ---------------------------------------------------------------- load
d = pd.read_csv(LEDGER, sep="\t", header=None, names=COLS, dtype=str)
doy = d.when_raw.str.extract(r"(\d+)(?:st|nd|rd|th)")[0].astype(int)
hms = d.when_raw.str.extract(r"(\d{2}:\d{2}:\d{2})")[0]
d["ts"] = pd.to_datetime("2026-01-01") + pd.to_timedelta(doy - 1, "D") + pd.to_timedelta(hms)
for c in ["pub", "prv", "auc_pub", "auc_prv", "f1_pub", "f1_prv"]:
    d[c] = d[c].astype(float)
d = d.sort_values("ts", ascending=False).reset_index(drop=True)
lb = pd.read_csv(BOARD, sep="\t")

# ---------------------------------------------------------------- page 1
doc.add_paragraph("Zindi Submission Ledger", style="Title")
para("GeoAI Aquaculture Pond Identification Challenge by FAO and ITU  |  captured 17 Aug 2026, "
     "after the private leaderboard was revealed", italic=True, color=GREY, space=12)

para("Final: rank 120 of 500 -- private 0.910686008. Winner 0.956900206.", bold=True, size=12)
para("Scored entry: submission_champion_dualpolmix10_regimematch.csv (one of the two designated "
     "finalists). 91 submissions across 87 distinct files, 7 Jul - 17 Aug 2026.", space=12)

para("How to read these tables", bold=True, size=11)
para("Leaderboard score = 0.6 x F1 + 0.4 x ROC-AUC, verified exactly on all 182 (AUC, F1, composite) "
     "triples in this document -- maximum residual 8e-10. The F1 column is computed at a hard 0.5 "
     "cut; the competition rules forbid choosing any other threshold.")
para("The public slice is 333 rows (181 positive) and the private slice 697 rows (379 positive); "
     "these were solved exactly, not assumed. Private scores were invisible until the competition "
     "closed. Rows shaded in yellow are the two entries designated as finalists.", space=14)

# ---------------------------------------------------------------- table 1
doc.add_paragraph("1. All 91 submissions", style="Heading 1")
para("Newest first. FILE is the exact CSV uploaded to Zindi.", italic=True, color=GREY, space=8)

rows = [(r.ts.strftime("%d %b %H:%M"), r.id, r.file, "FINAL" if r.sel == "SEL" else "",
         f"{r.pub:.9f}", f"{r.prv:.9f}", f"{r.auc_pub:.6f}", f"{r.auc_prv:.6f}",
         f"{r.f1_pub:.6f}", f"{r.f1_prv:.6f}") for _, r in d.iterrows()]
table(["DATE", "ID", "FILE", "", "PUBLIC", "PRIVATE", "AUC pub", "AUC prv", "F1 pub", "F1 prv"],
      rows,
      widths=[0.85, 0.72, 3.05, 0.45, 1.02, 1.02, 0.83, 0.83, 0.83, 0.83],
      highlight=lambda i, r: r[3] == "FINAL", mono_cols=(4, 5, 6, 7, 8, 9))

doc.add_page_break()

# ---------------------------------------------------------------- table 2
doc.add_paragraph("2. Final leaderboard -- top 25, and us", style="Heading 1")
para("500 teams finished. Ranks 1-75 were captured in full to "
     "experiments/zindi_final_leaderboard_top75.tsv; the top 25 are reproduced here.",
     italic=True, color=GREY, space=8)

sel = pd.concat([lb[lb["rank"] <= 25], lb[lb["rank"] == 120]])
rows = [(int(r["rank"]), r.user, int(r.n_sub), f"{r.pub:.9f}", f"{r.prv:.9f}",
         f"{r.auc_prv:.6f}", f"{r.f1_prv:.6f}") for _, r in sel.iterrows()]
table(["RANK", "TEAM", "SUBS", "PUBLIC", "PRIVATE", "AUC prv", "F1 prv"], rows,
      widths=[0.6, 2.4, 0.6, 1.3, 1.3, 1.1, 1.1],
      highlight=lambda i, r: r[0] == 120, mono_cols=(3, 4, 5, 6))

para("")
para("Ranks 3, 5, 6 and 7 posted byte-identical F1 on BOTH splits (public TP=173/PP=191, private "
     "TP=364/PP=397) while their AUC columns differ -- four independent teams shipping the same "
     "binary label set. A shared public approach existed.", italic=True, space=14)

# ---------------------------------------------------------------- table 3
doc.add_paragraph("3. Why we finished 120th", style="Heading 1")
para("Gap to first is 0.046214 = 0.013478 from AUC (29%) + 0.032736 from F1 (71%).", bold=True)
para("Below: every team whose private AUC is within 0.005 of ours -- an equally good global "
     "ranking -- with their private confusion cell recovered exactly by inverting the reported F1 "
     "on the solved private slice (n=697, P=379). Each inversion is unique.", space=8)

peers = [
    (49, "pmaurente", 0.945570, 0.921438, 346, 372, 0.9301, 0.9129, 0.5337),
    (45, "simonMakumi", 0.950192, 0.919598, 366, 417, 0.8777, 0.9657, 0.5983),
    (56, "tw_zent", 0.949038, 0.916230, 350, 385, 0.9091, 0.9235, 0.5524),
    (61, "Mutombwa", 0.952515, 0.908163, 356, 405, 0.8790, 0.9393, 0.5811),
    (71, "lesleygrin", 0.947313, 0.906702, 345, 382, 0.9031, 0.9103, 0.5481),
    (68, "mwarsssss", 0.950013, 0.906005, 347, 387, 0.8966, 0.9156, 0.5552),
    (75, "Nayal_17", 0.951801, 0.901660, 353, 404, 0.8738, 0.9314, 0.5796),
    (120, "forge (us)", 0.950158, 0.884371, 348, 408, 0.8529, 0.9182, 0.5854),
]
table(["RANK", "TEAM", "AUC prv", "F1 prv", "TP", "PP", "PRECISION", "RECALL", "POS-RATE"],
      [(a, b, f"{c:.6f}", f"{e:.6f}", f, g, f"{h:.4f}", f"{i:.4f}", f"{j:.4f}")
       for a, b, c, e, f, g, h, i, j in peers],
      widths=[0.6, 1.6, 1.05, 1.05, 0.6, 0.6, 1.0, 0.9, 0.95],
      highlight=lambda i, r: r[0] == 120, mono_cols=(2, 3, 6, 7, 8))

para("")
para("It is not where the 0.5 cut lands. Matching the operating point, Mutombwa predicts 3 FEWER "
     "positives and finds 8 more real ponds; Nayal_17 predicts 4 fewer and finds 5 more. Global AUC "
     "is identical, so their wrong pairs sit deep in the list where nothing is decided, and ours "
     "straddle the boundary. Our ranking is fine on average and weak exactly where the decision is "
     "made. True private prevalence is 379/697 = 0.5437; we ran at 0.5854.", space=12)

para("Full analysis: POST_MORTEM.md  |  reproduce every number: python tools/post_mortem.py",
     italic=True, color=GREY)

doc.save(OUT)
print("saved:", OUT)
